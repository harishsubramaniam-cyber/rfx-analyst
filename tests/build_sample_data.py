"""Fabricate the five supplier responses.

The brief asks for a dataset a procurement person would nod at, so the point of
this file is not to produce five tidy files with a few planted traps. It is to
produce five documents that look like they came out of five different offices:

  * prices derived from a cost model with per-supplier strengths, so nobody
    wins everything and no column is a fixed multiple of another
  * each supplier's own item codes, own column order, own house wording
  * GSTIN, HSN codes, quote references, validity dates, MOQs, slab rates,
    contact blocks, revision notes -- the furniture real quotations carry
  * the difficulties arriving the way they arrive in real life: a rate card
    photographed on a phone, an exporter quoting FOB in dollars, a discount
    in eight-point type at the bottom of page two

Run:  python tests/build_sample_data.py
"""

from __future__ import annotations

import os
import random
import sys
from dataclasses import dataclass
from email.message import EmailMessage

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
sys.path.insert(0, ROOT)

from core import brand  # noqa: E402

OUT = os.path.join(ROOT, "sample_data")
ATTACH = os.path.join(OUT, "attachments")

RNG = random.Random(20260827)

USD_INR = 87.50          # the rate the buyer's finance team publishes
QUOTE_DATE = "27 August 2026"
PO_DATE = "12 September 2025"    # last year: the PO a buyer starts from


# ---------------------------------------------------------------------------
# the catalogue the buyer sent out
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class Item:
    sku: str
    description: str
    quantity: int
    base: float          # indicative market cost, INR per piece
    family: str          # box | printed | heavy | diecut | spec | partition
                         # | sheet | roll | protector | consumable
    moq: int


CATALOGUE: list[Item] = [
    Item("BX-001", "5-ply corrugated box 400x300x250 mm", 10000, 44.00, "box", 500),
    Item("BX-002", "5-ply corrugated box 450x350x300 mm", 8000, 51.00, "box", 500),
    Item("BX-003", "5-ply corrugated box 500x400x350 mm", 7000, 58.00, "box", 500),
    Item("BX-004", "5-ply corrugated box 600x400x400 mm", 6000, 75.00, "box", 250),
    Item("BX-005", "3-ply corrugated box 300x250x200 mm", 12000, 25.00, "box", 1000),
    Item("BX-006", "3-ply corrugated box 350x250x200 mm", 10000, 29.00, "box", 1000),
    Item("BX-007", "3-ply corrugated box 400x300x250 mm", 9000, 36.00, "box", 1000),
    Item("BX-008", "3-ply corrugated box 450x300x300 mm", 7000, 43.00, "box", 500),
    Item("BX-009", "Printed shipper box 400x300x250 mm", 5000, 61.00, "printed", 1000),
    Item("BX-010", "Printed shipper box 500x400x300 mm", 4000, 76.00, "printed", 1000),
    Item("BX-011", "Heavy-duty 7-ply box 600x500x500 mm", 2500, 131.00, "heavy", 250),
    Item("BX-012", "Heavy-duty 7-ply box 800x600x500 mm", 1800, 172.00, "heavy", 250),
    Item("BX-013", "Die-cut box 300x200x150 mm", 6000, 33.00, "diecut", 500),
    Item("BX-014", "Die-cut box 400x300x200 mm", 5000, 41.00, "diecut", 500),
    Item("BX-015", "Die-cut box 500x350x250 mm", 4000, 54.00, "diecut", 500),
    Item("BX-016", "5-ply box, moisture-resistant kraft", 4500, 71.00, "spec", 500),
    Item("BX-017", "5-ply box, burst strength 16 kg/cm2", 3500, 77.00, "spec", 500),
    Item("BX-018", "5-ply box, burst strength 18 kg/cm2", 3000, 86.00, "spec", 500),
    Item("BX-019", "3-ply box, burst strength 12 kg/cm2", 8000, 30.00, "spec", 1000),
    Item("BX-020", "3-ply box, burst strength 14 kg/cm2", 6500, 38.00, "spec", 1000),
    Item("BX-021", "Partition set, 12-cell", 5000, 19.00, "partition", 1000),
    Item("BX-022", "Partition set, 24-cell", 3500, 26.00, "partition", 1000),
    Item("BX-023", "Corrugated sheet 1000x800 mm", 9000, 64.00, "sheet", 1000),
    Item("BX-024", "Corrugated sheet 1200x1000 mm", 7000, 87.00, "sheet", 1000),
    Item("BX-025", "5-ply roll, 1200 mm width", 1200, 430.00, "roll", 50),
    Item("BX-026", "Edge protector 50x50x5 mm", 15000, 12.50, "protector", 2000),
    Item("BX-027", "Corner protector 75x75x5 mm", 10000, 17.50, "protector", 2000),
    Item("BX-028", "Kraft paper tape 48 mm x 50 m", 20000, 39.00, "consumable", 1000),
    Item("BX-029", "Stretch film 500 mm x 300 m", 6000, 191.00, "consumable", 200),
    Item("BX-030", "Pallet top sheet 1200x1000 mm", 5000, 29.00, "sheet", 1000),
]

BY_SKU = {item.sku: item for item in CATALOGUE}

# HSN codes a real quotation would carry
HSN = {"box": "48191010", "printed": "48191010", "heavy": "48191010",
       "diecut": "48192090", "spec": "48191010", "partition": "48195000",
       "sheet": "48081000", "roll": "48081000", "protector": "48239019",
       "consumable": "48234000"}


# ---------------------------------------------------------------------------
# pricing: each supplier is genuinely better at some things
# ---------------------------------------------------------------------------

STRENGTHS: dict[str, dict[str, float]] = {
    # multiplier by family; < 1 means competitive
    "shakti": {"box": 0.96, "printed": 1.06, "heavy": 1.09, "diecut": 1.02,
               "spec": 1.00, "partition": 0.94, "sheet": 1.03, "roll": 1.05,
               "protector": 0.97, "consumable": 1.01},
    "balaji": {"box": 1.02, "printed": 0.93, "heavy": 1.04, "diecut": 0.95,
               "spec": 1.01, "partition": 1.03, "sheet": 1.06, "roll": 1.08,
               "protector": 1.05, "consumable": 1.02},
    "meridian": {"box": 1.05, "printed": 1.02, "heavy": 1.01, "diecut": 1.06,
                 "spec": 1.04, "partition": 1.05, "sheet": 0.95, "roll": 0.93,
                 "protector": 1.04, "consumable": 0.97},
    "northstar": {"box": 1.01, "printed": 1.04, "heavy": 0.92, "diecut": 1.03,
                  "spec": 0.97, "partition": 1.00, "sheet": 1.02, "roll": 1.01,
                  "protector": 0.91, "consumable": 1.06},
    "pacific": {"box": 1.08, "printed": 1.01, "heavy": 0.96, "diecut": 1.04,
                "spec": 0.99, "partition": 1.09, "sheet": 1.01, "roll": 0.98,
                "protector": 1.12, "consumable": 0.94},
}


def price_for(vendor: str, item: Item) -> float:
    """Cost model plus supplier strength plus a little quoting noise."""
    factor = STRENGTHS[vendor][item.family]
    noise = RNG.gauss(0, 0.018)
    value = item.base * factor * (1 + noise)
    # quotations round the way humans round
    if value >= 100:
        return round(value, 0) if RNG.random() < 0.45 else round(value, 2)
    if value >= 20:
        return round(value * 4) / 4 if RNG.random() < 0.3 else round(value, 2)
    return round(value, 2)


PRICES: dict[str, dict[str, float]] = {
    vendor: {item.sku: price_for(vendor, item) for item in CATALOGUE}
    for vendor in STRENGTHS
}


def rupees(value: float) -> str:
    return f"{value:,.2f}"


# ---------------------------------------------------------------------------
# 1 · Shakti Packaging -- spreadsheet in the supplier's own house format
# ---------------------------------------------------------------------------

def build_shakti() -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Quotation"

    thin = Side(style="thin", color="BFBFBF")
    box = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_fill = PatternFill("solid", fgColor="1F3864")
    head_font = Font(bold=True, color="FFFFFF", size=9)

    # --- letterhead block, merged, nothing like the buyer's template --------
    ws.merge_cells("A1:I1")
    ws["A1"] = "SHAKTI PACKAGING INDUSTRIES PVT LTD"
    ws["A1"].font = Font(bold=True, size=15, color="1F3864")
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A2:I2")
    ws["A2"] = ("Plot 44, Phase II, Peenya Industrial Area, Bengaluru 560058  |  "
                "GSTIN 29AAGCS4471P1ZK  |  sales@shaktipack.co.in  |  +91 80 2839 4410")
    ws["A2"].font = Font(size=8, color="595959")
    ws["A2"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A4:I4")
    ws["A4"] = "QUOTATION — CORRUGATED PACKAGING (ANNUAL RATE CONTRACT)"
    ws["A4"].font = Font(bold=True, size=11)
    ws["A4"].alignment = Alignment(horizontal="center")

    for row, (label, value) in enumerate(
        [("Quotation No.", "QT/2026-27/0812"),
         ("Date", QUOTE_DATE),
         ("Buyer", brand.COMPANY_LEGAL),
         ("Kind Attn.", brand.CONTACT),
         ("Validity", "45 days from date of quotation"),
         ("Ref.", "Your enquiry RFX/CORR/2026-08 dated 18 Aug 2026")],
        start=6,
    ):
        ws.cell(row=row, column=1, value=label).font = Font(bold=True, size=9)
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=5)
        ws.cell(row=row, column=2, value=value).font = Font(size=9)

    # --- their column order, not the buyer's -------------------------------
    header_row = 12
    headers = ["Sr", "Item Code", "Description", "HSN", "UOM",
               "MOQ", "Rate (INR)", "Remarks"]
    for column, title in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=column, value=title)
        cell.fill, cell.font, cell.border = head_fill, head_font, box
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Shakti does not quote the two lightest 3-ply lines this year.
    skipped = {"BX-019", "BX-005"}
    remarks = {
        "BX-004": "Rate for kraft 180 GSM outer; 200 GSM +₹2.10",
        "BX-011": "Sub-contracted; lead time 30 days",
        "BX-016": "Poly-coated liner, moisture barrier as per spec",
        "BX-025": "Mill roll — supplied in 50 kg reels",
        "BX-029": "Imported film, subject to FX at time of despatch",
    }

    row = header_row + 1
    serial = 1
    for item in CATALOGUE:
        if item.sku in skipped:
            continue
        # the buyer's own item code, written loosely by hand on some rows
        buyer_ref = item.sku if serial % 4 else item.sku.replace("-", " ")
        values = [
            serial, buyer_ref, item.description, HSN[item.family],
            "Nos" if item.family not in {"sheet", "roll", "consumable"} else
            ("Sheet" if item.family == "sheet" else "Roll"),
            item.moq, PRICES["shakti"][item.sku], remarks.get(item.sku, ""),
        ]
        for column, value in enumerate(values, start=1):
            cell = ws.cell(row=row, column=column, value=value)
            cell.border = box
            cell.font = Font(size=9)
            if column == 7:
                cell.number_format = "#,##0.00"
                cell.alignment = Alignment(horizontal="right")
        row += 1
        serial += 1

    # --- terms, below the table, in the supplier's own words ---------------
    row += 1
    terms = [
        "TERMS & CONDITIONS",
        "1.  Rates are Ex-Works Bengaluru. Freight to your Bommasandra warehouse is included "
        "for despatches above 2 tonnes; below that, freight at actuals.",
        "2.  GST @ 18% extra as applicable.",
        "3.  Payment: 30 days from date of invoice.",
        "4.  Delivery: 18–21 days from receipt of firm PO and artwork approval.",
        "5.  Rates are based on kraft paper at ₹42/kg. Any variation beyond ±5% will be "
        "passed on with 30 days' notice.",
        "6.  Slab rates for higher offtake are given in the 'Slab Rates' sheet.",
        "7.  ISO 9001:2015 certificate enclosed (Shakti_ISO9001_certificate.pdf).",
    ]
    for line in terms:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
        cell = ws.cell(row=row, column=1, value=line)
        cell.font = Font(bold=line.startswith("TERMS"), size=9)
        cell.alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[row].height = 26 if len(line) > 90 else 14
        row += 1

    row += 2
    ws.cell(row=row, column=1, value="For SHAKTI PACKAGING INDUSTRIES PVT LTD").font = \
        Font(bold=True, size=9)
    ws.cell(row=row + 2, column=1, value="R. Venkatesh").font = Font(size=9)
    ws.cell(row=row + 3, column=1, value="Sr. Manager — Sales").font = Font(size=8)

    widths = [4, 11, 42, 10, 7, 8, 12, 38]
    for column, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(column)].width = width
    ws.freeze_panes = "A13"

    # --- a second sheet with volume breaks ---------------------------------
    slab = wb.create_sheet("Slab Rates")
    slab["A1"] = "VOLUME SLAB RATES (applicable on annual offtake per item)"
    slab["A1"].font = Font(bold=True, size=11)
    for column, title in enumerate(
            ["Item Code", "Up to 5,000", "5,001 – 15,000", "Above 15,000"],
            start=1):
        cell = slab.cell(row=3, column=column, value=title)
        cell.font, cell.fill, cell.border = head_font, head_fill, box
    slab_row = 4
    for item in CATALOGUE[:12]:
        if item.sku in skipped:
            continue
        base = PRICES["shakti"][item.sku]
        for column, value in enumerate(
                [item.sku,
                 round(base, 2), round(base * 0.975, 2), round(base * 0.955, 2)],
                start=1):
            cell = slab.cell(row=slab_row, column=column, value=value)
            cell.border, cell.font = box, Font(size=9)
            if column >= 2:
                cell.number_format = "#,##0.00"
        slab_row += 1
    slab.cell(row=slab_row + 1, column=1,
              value="Slab rates apply on confirmed annual volumes only.").font = Font(size=8)
    for column, width in enumerate([12, 14, 16, 15], start=1):
        slab.column_dimensions[get_column_letter(column)].width = width

    # --- questionnaire, answered in their own phrasing ---------------------
    q = wb.create_sheet("Vendor Declaration")
    q["A1"] = "VENDOR QUALITY DECLARATION"
    q["A1"].font = Font(bold=True, size=11)
    rows = [
        ("ISO 9001 certification held?", "Yes — valid to Nov 2027, copy enclosed"),
        ("Monthly production capacity above 50,000 units?", "Yes — approx. 90,000 units/month"),
        ("On-time delivery performance above 95%?", "Yes — 96.4% for FY 2025-26"),
        ("Quality rejection rate below 1%?", "Yes — 0.6% average"),
        ("Able to commit to a 30-day delivery SLA?", "Yes"),
    ]
    q.cell(row=3, column=1, value="Question").font = Font(bold=True, size=9)
    q.cell(row=3, column=2, value="Response").font = Font(bold=True, size=9)
    for index, (question, answer) in enumerate(rows, start=4):
        q.cell(row=index, column=1, value=question).font = Font(size=9)
        q.cell(row=index, column=2, value=answer).font = Font(size=9)
    q.column_dimensions["A"].width = 48
    q.column_dimensions["B"].width = 46

    wb.save(os.path.join(OUT, "Shakti_Packaging_Quotation.xlsx"))


# ---------------------------------------------------------------------------
# 2 · Sri Balaji Corrugators -- PDF on letterhead, discount buried at the end
# ---------------------------------------------------------------------------

def build_balaji() -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)

    styles = getSampleStyleSheet()
    small = ParagraphStyle("small", parent=styles["Normal"], fontSize=7.5, leading=9.5)
    tiny = ParagraphStyle("tiny", parent=styles["Normal"], fontSize=6.2, leading=8,
                          textColor=colors.HexColor("#444444"))
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=8.5, leading=11.5)
    head = ParagraphStyle("head", parent=styles["Normal"], fontSize=16,
                          leading=19, textColor=colors.HexColor("#7A1F1F"),
                          spaceAfter=1)

    story = []
    story.append(Paragraph("SRI BALAJI CORRUGATORS PVT LTD", head))
    story.append(Paragraph(
        "No. 118/2, Hosur Road, Bommasandra Industrial Area, Bengaluru 560099<br/>"
        "GSTIN 29AACCS9921L1Z4 &nbsp;|&nbsp; CIN U21029KA2009PTC050113 &nbsp;|&nbsp; "
        "balaji.sales@sbcorrugators.in &nbsp;|&nbsp; +91 80 4123 7788", small))
    story.append(Spacer(1, 7 * mm))

    story.append(Paragraph("<b>COMMERCIAL QUOTATION</b>", body))
    story.append(Paragraph(
        f"Quote Ref: SBC/RFX/0826-R2 &nbsp;&nbsp;|&nbsp;&nbsp; Date: {QUOTE_DATE}"
        f"<br/>To: {brand.COMPANY_LEGAL}, {brand.ADDRESS}<br/>"
        "Sub: Supply of corrugated packaging — annual rate contract 2026-27<br/>"
        "<i>This revision R2 supersedes our quotation SBC/RFX/0826-R1 dated "
        "14 Aug 2026.</i>", small))
    story.append(Spacer(1, 5 * mm))

    # Balaji does not quote the two heaviest lines; they sub-contract those.
    skipped = {"BX-012", "BX-025"}
    substitutions = {
        "BX-016": "Offered in poly-laminated kraft in place of specified wax coating",
        "BX-018": "BF-18 grade; BF-20 available at +4%",
    }

    data = [["Sl", "Item Code", "Description", "HSN", "MOQ", "Rate", "Unit"]]
    serial = 1
    for item in CATALOGUE:
        if item.sku in skipped:
            continue
        unit = {"sheet": "INR / sheet", "roll": "INR / roll",
                "consumable": "INR / roll"}.get(item.family, "INR / box")
        if item.family == "partition":
            unit = "INR / set"
        description = item.description
        if item.sku in substitutions:
            description += f"<br/><font size=6 color='#7A1F1F'>{substitutions[item.sku]}</font>"
        data.append([
            str(serial), item.sku, Paragraph(description, small), HSN[item.family],
            f"{item.moq:,}", f"₹{rupees(PRICES['balaji'][item.sku])}", unit,
        ])
        serial += 1

    table = Table(data, colWidths=[9 * mm, 20 * mm, 62 * mm, 18 * mm, 16 * mm,
                                   22 * mm, 22 * mm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#7A1F1F")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#BBBBBB")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (5, 1), (5, -1), "RIGHT"),
        ("ALIGN", (4, 1), (4, -1), "RIGHT"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F7F3F3")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 2.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
    ]))
    story.append(table)
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(
        "BX-012 and BX-025 are not quoted in this revision — these are outside our "
        "in-house capability and we would rather not sub-contract them.", small))

    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph("<b>Commercial Terms</b>", body))
    story.append(Paragraph(
        "Freight is included for deliveries to your Bommasandra warehouse. "
        "GST @ 18% extra. Payment terms 45 days from invoice. Rates firm for 60 days "
        "from the date of this quotation. Delivery 20–24 days from firm order; "
        "printed items require 5 additional working days for plate approval.", small))

    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph("<b>Quality Questionnaire</b>", body))
    story.append(Paragraph(
        "ISO 9001 certified: <b>Yes</b> (certificate no. IN-QMS-44192, enclosed)<br/>"
        "Monthly production capacity &gt; 50,000 units: <b>Yes</b><br/>"
        "On-time delivery &gt; 95%: <b>Yes</b> — 97.1% in FY 2025-26<br/>"
        "Quality rejection rate &lt; 1%: <b>Yes</b> — 0.4%<br/>"
        "Can meet 30-day delivery SLA: <b>Yes</b>", small))

    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph(
        "Enclosures: (1) ISO 9001:2015 certificate  (2) Burst strength test report, "
        "SGS Bengaluru, dated 02 Aug 2026  (3) Company profile", small))

    # the discount, in eight-point type, at the very bottom of the last page
    story.append(Spacer(1, 22 * mm))
    story.append(Paragraph(
        "Notes: (a) Rates quoted are for kraft of 180 GSM unless otherwise stated. "
        "(b) Artwork charges for printed cartons are one-time and are not included "
        "above. (c) <b>An additional discount of 2% on the quoted rates will apply to "
        "the entire order value where annualised offtake across all items exceeds "
        "50,000 units.</b> (d) Any statutory levy introduced after the date of this "
        "quotation will be charged at actuals. (e) This quotation is confidential.",
        tiny))

    doc = SimpleDocTemplate(
        os.path.join(OUT, "Sri_Balaji_Corrugators_Quote.pdf"), pagesize=A4,
        leftMargin=16 * mm, rightMargin=16 * mm, topMargin=15 * mm, bottomMargin=14 * mm,
        title="Sri Balaji Corrugators — Quotation SBC/RFX/0826-R2",
    )
    doc.build(story)


# ---------------------------------------------------------------------------
# 3 · Meridian Packaging -- Word quotation, commercials written as prose
# ---------------------------------------------------------------------------

def build_meridian() -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    title = document.add_paragraph()
    run = title.add_run("MERIDIAN PACKAGING LLP")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1B, 0x4D, 0x3E)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = document.add_paragraph()
    sub_run = sub.add_run(
        "Survey 71/3, Jigani Link Road, Anekal Taluk, Bengaluru 562106\n"
        "GSTIN 29ABLFM7712R1ZQ  |  accounts@meridianpack.in  |  +91 99866 21140")
    sub_run.font.size = Pt(8)
    sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    # A quotation document, not correspondence: no attention line, no subject,
    # no salutation and no sign-off. What makes this response hard to read is
    # that the commercials are written as prose with no table anywhere -- that
    # is the whole reason this sample exists, and it is untouched. The letter
    # wrapper was never part of the difficulty, only of the dressing.
    doc_title = document.add_paragraph()
    doc_title_run = doc_title.add_run("QUOTATION")
    doc_title_run.bold = True
    doc_title_run.font.size = Pt(13)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        f"Quotation No: MPL/COMM/2026/311{' ' * 10}Date: {QUOTE_DATE}")

    recipient = document.add_paragraph()
    recipient.add_run("For: ")
    recipient.add_run(brand.COMPANY_LEGAL).bold = True
    document.add_paragraph(brand.ADDRESS)
    document.add_paragraph(
        "Against enquiry: RFX/CORR/2026-08 — corrugated packaging, FY 2026-27")

    document.add_paragraph()
    basis = document.add_paragraph()
    basis_run = basis.add_run("Scope and Basis of Pricing")
    basis_run.bold = True
    basis_run.font.size = Pt(12)

    document.add_paragraph(
        "Rates are quoted against the specifications shared with the enquiry. Where "
        "an item is supplied as a set rather than an individual carton, the rate is "
        "per set; sheet and roll products are priced on their normal selling unit. "
        "Rates assume the annual volumes indicated in the enquiry and are subject to "
        "the minimum order quantities noted against each line.")

    document.add_paragraph()
    heading = document.add_paragraph()
    heading_run = heading.add_run("Commercial Offer")
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    # Meridian prices everything, but writes it as prose -- no table anywhere.
    for item in CATALOGUE:
        price = PRICES["meridian"][item.sku]
        unit = {"sheet": "per sheet", "roll": "per roll",
                "consumable": "per roll"}.get(item.family, "per box/set")
        if item.family == "protector":
            unit = "per box/set"
        lead = 24 if item.family in {"box", "spec"} else 28
        sentence = (
            f"For {item.sku}, {item.description}, against your indicated annual "
            f"quantity of {item.quantity:,}, our rate is ₹{rupees(price)} {unit}, "
            f"minimum order {item.moq:,}. Delivery approximately {lead} days from "
            f"receipt of order."
        )
        if item.sku == "BX-009":
            sentence += (" Printing plates are a one-time charge and are quoted "
                         "separately on approval of artwork.")
        if item.sku == "BX-029":
            sentence += (" This is a traded item and the rate moves with the "
                         "imported film market; we will reconfirm at the time of order.")
        document.add_paragraph(sentence)

    conditions = document.add_paragraph()
    conditions_run = conditions.add_run("Commercial Conditions")
    conditions_run.bold = True
    conditions_run.font.size = Pt(12)

    document.add_paragraph(
        "Freight is charged separately at actuals and is not included in the rates "
        "above; we can arrange despatch through your nominated transporter if you "
        "prefer. GST at 18% is extra. Prices hold for 30 days from the date of this "
        "quotation. Payment terms requested are 30 days from invoice. A volume rebate of "
        "1.5% is applicable where cumulative purchases in any quarter exceed ₹50 lakh, "
        "settled by credit note at quarter end.")

    quality = document.add_paragraph()
    quality_run = quality.add_run("Quality Questionnaire")
    quality_run.bold = True
    quality_run.font.size = Pt(12)

    document.add_paragraph("ISO 9001 certified: Yes (certificate valid to March 2028)")
    document.add_paragraph("Monthly production capacity above 50,000 units: Yes")
    document.add_paragraph(
        "On-time delivery above 95%: No — our measured performance for the last "
        "financial year was 91.8%. We have added a second shift in the flexo line and "
        "expect to be above 95% from Q3.")
    document.add_paragraph("Quality rejection rate below 1%: Yes (0.7%)")
    document.add_paragraph("Can meet a 30-day delivery SLA: Yes")

    document.add_paragraph()
    document.add_paragraph(
        "Our ISO certificate and the last two third-party test reports are attached as "
        "Annexure A. We would be glad to arrange a plant visit at your convenience.")
    document.add_paragraph()
    document.add_paragraph("For Meridian Packaging LLP")
    document.add_paragraph()
    signature = document.add_paragraph()
    signature_run = signature.add_run("Kavitha Menon")
    signature_run.bold = True
    document.add_paragraph("Partner · Authorised Signatory")

    document.save(os.path.join(OUT, "Meridian_Packaging_Offer.docx"))


# ---------------------------------------------------------------------------
# 4 · Northstar -- a printed rate card, photographed on a phone
# ---------------------------------------------------------------------------

def _render_rate_card() -> "PIL.Image.Image":
    from PIL import Image, ImageDraw, ImageFont

    width, height = 1654, 2339          # A4 at 200 dpi
    card = Image.new("RGB", (width, height), (252, 251, 247))
    draw = ImageDraw.Draw(card)

    def font(size: int, bold: bool = False):
        names = (["DejaVuSans-Bold.ttf", "DejaVuSans.ttf"] if bold
                 else ["DejaVuSans.ttf"])
        for name in names:
            for base in ("/usr/share/fonts/truetype/dejavu/",
                         "/usr/share/fonts/truetype/"):
                try:
                    return ImageFont.truetype(base + name, size)
                except OSError:
                    continue
        return ImageFont.load_default()

    x0 = 120
    draw.text((x0, 120), "NORTHSTAR PACKAGING", font=font(52, True), fill=(20, 20, 20))
    draw.text((x0, 190), "Unit 7, KIADB Industrial Estate, Hoskote, Bengaluru 562114",
              font=font(24), fill=(90, 90, 90))
    draw.text((x0, 226), "GSTIN 29AAFCN2288K1ZP   Ph 080-2797 6612",
              font=font(24), fill=(90, 90, 90))
    draw.line((x0, 268, width - x0, 268), fill=(40, 40, 40), width=3)

    draw.text((x0, 300), "RATE CARD  II  —  CORRUGATED PRODUCTS",
              font=font(34, True), fill=(20, 20, 20))
    draw.text((x0, 348), "w.e.f. 01-08-2026   |   Supersedes Rate Card I",
              font=font(24), fill=(90, 90, 90))
    draw.text((x0, 350), f"Prepared for: {brand.COMPANY_LEGAL}, {brand.ADDRESS}",
              font=font(23), fill=(70, 70, 70))
    draw.text((x0, 384), "All rates subject to confirmation at time of order. Freight extra.",
              font=font(23), fill=(90, 90, 90))

    # Northstar does not list two of the buyer's items at all
    skipped = {"BX-008", "BX-019"}
    # and prices a few things on their own basis
    # Northstar prices a few things on their own basis. The bulk figure is the
    # same underlying rate, expressed per hundred -- so it reconciles against
    # the other quotes once converted, which is the point.
    special = {
        "BX-009": (f"{PRICES['northstar']['BX-009'] * 100:,.0f}", "/ 100 pcs"),
        "BX-016": ("68", "/ kg"),
        "BX-025": (f"{PRICES['northstar']['BX-025']:,.0f}", "/ roll"),
        "BX-001": ("", "/ box"),
        "BX-013": ("", "/box"),
    }

    rate_x = x0 + 820           # keep the span short so perspective cannot
                                # drift a rate away from its own row
    y = 452
    draw.text((x0, y), "CODE", font=font(22, True), fill=(70, 70, 70))
    draw.text((x0 + 170, y), "DESCRIPTION", font=font(22, True), fill=(70, 70, 70))
    draw.text((rate_x, y), "RATE", font=font(22, True), fill=(70, 70, 70))
    y += 34
    draw.line((x0, y, width - x0, y), fill=(150, 150, 150), width=2)
    y += 18

    row_index = 0
    for item in CATALOGUE:
        if item.sku in skipped:
            continue
        price = PRICES["northstar"][item.sku]
        if item.sku in special and special[item.sku][0]:
            shown, suffix = special[item.sku]
        elif item.sku in special:
            shown, suffix = f"{price:,.0f}", special[item.sku][1]
        else:
            shown, suffix = (f"{price:,.0f}" if price >= 100 else f"{price:,.2f}"), ""

        # alternating band + hairline rule: this is how printed rate cards keep
        # the eye on one row, and it makes the photograph unambiguous
        if row_index % 2 == 0:
            draw.rectangle((x0 - 12, y - 9, width - x0 + 12, y + 37),
                           fill=(240, 239, 232))
        draw.line((x0 - 12, y + 37, width - x0 + 12, y + 37),
                  fill=(214, 212, 205), width=1)

        draw.text((x0, y), item.sku, font=font(24), fill=(25, 25, 25))
        draw.text((x0 + 170, y), item.description[:52], font=font(24), fill=(25, 25, 25))
        draw.text((rate_x, y), f"{shown} {suffix}".strip(), font=font(24),
                  fill=(25, 25, 25))
        row_index += 1
        y += 46

    y += 26
    draw.line((x0, y, width - x0, y), fill=(150, 150, 150), width=2)
    y += 26
    draw.text((x0, y), "Quality:  ISO 9001 - No (audit scheduled Nov 2026)",
              font=font(23), fill=(40, 40, 40)); y += 38
    draw.text((x0, y), "Capacity: Yes, 60,000 units/month     OTD >95%: Yes",
              font=font(23), fill=(40, 40, 40)); y += 38
    draw.text((x0, y), "Note: BX-008 and BX-019 are not manufactured by us.",
              font=font(23), fill=(40, 40, 40)); y += 60
    draw.text((x0, y), "MOQ 500 nos per item unless agreed otherwise.",
              font=font(23), fill=(90, 90, 90))

    # a handwritten-looking annotation, as scribbled by whoever sent it
    draw.text((x0 + 60, y + 90), "rev.2 — confirm rates w/ Suresh before PO",
              font=font(30), fill=(24, 44, 150))

    return card


def build_northstar() -> None:
    """Render the card, then make it look like a phone photo of paper."""
    import numpy as np
    from PIL import Image, ImageEnhance, ImageFilter

    card = _render_rate_card()
    width, height = card.size

    # --- perspective: the phone was not held square to the page ------------
    scene_w, scene_h = 1500, 2000
    corners = [(214, 118), (1288, 118), (1376, 1884), (126, 1884)]

    def perspective_coefficients(target, source):
        matrix = []
        for (tx, ty), (sx, sy) in zip(target, source):
            matrix.append([tx, ty, 1, 0, 0, 0, -sx * tx, -sx * ty])
            matrix.append([0, 0, 0, tx, ty, 1, -sy * tx, -sy * ty])
        A = np.matrix(matrix, dtype=float)
        B = np.array(source).reshape(8)
        return np.array(np.dot(np.linalg.inv(A.T * A) * A.T, B)).reshape(8)

    coefficients = perspective_coefficients(
        corners, [(0, 0), (width, 0), (width, height), (0, height)])

    desk = Image.new("RGB", (scene_w, scene_h), (86, 78, 70))
    # a bit of desk texture
    noise = (np.random.default_rng(7).normal(0, 7, (scene_h, scene_w, 3))
             + np.array([86, 78, 70]))
    desk = Image.fromarray(np.clip(noise, 0, 255).astype("uint8"))

    warped = card.transform((scene_w, scene_h), Image.PERSPECTIVE, coefficients,
                            Image.BICUBIC)
    mask = Image.new("L", (width, height), 255).transform(
        (scene_w, scene_h), Image.PERSPECTIVE, coefficients, Image.BICUBIC)

    # soft drop shadow under the sheet
    shadow = mask.filter(ImageFilter.GaussianBlur(18)).point(lambda v: int(v * 0.55))
    desk.paste((30, 26, 22), (6, 10), shadow)
    desk.paste(warped, (0, 0), mask)

    photo = np.asarray(desk).astype(np.float32)

    # --- lighting: window on the left, phone shadow on the right ----------
    yy, xx = np.mgrid[0:scene_h, 0:scene_w]
    gradient = (1.06 - 0.34 * (xx / scene_w) - 0.10 * (yy / scene_h))
    vignette = 1 - 0.30 * (((xx - scene_w / 2) / (scene_w / 2)) ** 2
                           + ((yy - scene_h / 2) / (scene_h / 2)) ** 2)
    photo *= (gradient * vignette)[:, :, None]

    # --- sensor noise and a warm indoor white balance ---------------------
    rng = np.random.default_rng(11)
    photo += rng.normal(0, 3.4, photo.shape)
    photo *= np.array([1.035, 1.0, 0.955])
    photo = np.clip(photo, 0, 255).astype("uint8")

    result = Image.fromarray(photo)
    result = result.filter(ImageFilter.GaussianBlur(0.7))       # hand shake
    result = ImageEnhance.Contrast(result).enhance(0.94)
    result = result.rotate(-0.7, resample=Image.BICUBIC, expand=False,
                           fillcolor=(78, 71, 64))
    result = result.resize((1200, 1600), Image.LANCZOS)

    result.save(os.path.join(OUT, "Northstar_RateCard_photo.jpg"),
                quality=72, optimize=True)


# ---------------------------------------------------------------------------
# 5 · Pacific Pack Global -- an exporter's email, quoting FOB in US dollars
# ---------------------------------------------------------------------------

def build_pacific() -> None:
    # not quoted: three items they do not carry
    skipped = {"BX-012", "BX-024"}
    lines = []
    for item in CATALOGUE:
        if item.sku in skipped:
            continue
        inr = PRICES["pacific"][item.sku]
        usd = inr / USD_INR

        if item.sku == "BX-005":
            lines.append(f"{item.sku} - USD {usd * 1000:,.0f} per 1000 pcs")
        elif item.sku == "BX-001":
            lines.append(f"{item.sku} - USD {usd / 0.42:,.2f} per kg "
                         "(we sell this grade by weight)")
        elif item.sku == "BX-009":
            lines.append(f"{item.sku} - USD {usd:,.2f} each, plus one-time "
                         "plate charge to be advised")
        elif item.sku == "BX-017":
            lines.append(f"{item.sku} - USD {usd:,.2f}, balance grades same as "
                         "our 2025 contract rates")
        elif item.sku == "BX-029":
            lines.append(f"{item.sku} - USD {usd * 100:,.0f} per 100 rolls")
        elif item.family == "sheet":
            lines.append(f"{item.sku} - USD {usd:,.2f} per sheet")
        elif item.family in {"roll", "consumable"}:
            lines.append(f"{item.sku} - USD {usd:,.2f} per roll")
        else:
            lines.append(f"{item.sku} - USD {usd:,.2f} per box")

    body = f"""Hi {brand.CONTACT.split(",")[0].split()[0]},

Apologies for the delay, I was travelling. Please find our offer below.

Note we invoice in USD as we are an export-oriented unit; conversion at your
banker's TT rate on invoice date. Prices are FOB Chennai. GST/IGST as
applicable extra.

{chr(10).join(lines)}

For anything not listed above please refer to our standard 2025 rate card
which we shared with your team in October last year - those rates still hold.

Lead time 30-35 days ex-Chennai. Payment 60 days from BL date.
MOQ 1000 pcs per line item, 500 for the 7-ply.

Quality questionnaire:
- ISO 9001 certified: Yes
- Monthly production capacity > 50,000 units: No (we run at about 35,000)
- On-time delivery > 95%: Yes
- Quality rejection rate < 1%: Yes
- Can meet 30-day delivery SLA: No - please see lead time above

Happy to discuss. I am in Bengaluru next Tuesday if you would like to meet.

Best regards,
Rajesh Pillai
Head - Exports | Pacific Pack Global Pvt Ltd
SIPCOT Industrial Park, Sriperumbudur, Tamil Nadu 602105
M +91 98404 33121

> On 18 Aug 2026, at 10:14, {brand.CONTACT.split(",")[0]} wrote:
> Dear Supplier, please find attached our enquiry for corrugated packaging
> covering 30 line items for FY 2026-27. Kindly revert with your best rates
> along with the completed quality questionnaire by 25 August.
"""

    # A real .eml rather than a .txt transcript. A saved message is what a
    # buyer actually forwards out of Outlook, and it exercises the reader that
    # has to cope with headers, encodings and attachments -- the text version
    # only ever exercised "read the file".
    message = EmailMessage()
    message["From"] = "Rajesh Pillai <rajesh.p@pacificpackglobal.com>"
    message["To"] = f"{brand.CONTACT_EMAIL}"
    message["Cc"] = "exports@pacificpackglobal.com"
    message["Date"] = "Thu, 27 Aug 2026 18:42:10 +0530"
    message["Subject"] = "RE: RFX/CORR/2026-08 - Corrugated packaging enquiry"
    message["Message-ID"] = "<2026082718421.rajesh@pacificpackglobal.com>"
    # Quoted-printable is what a mail client would do with a rupee sign or a
    # long line, and it is exactly what naive text reading gets wrong.
    message.set_content(body, cte="quoted-printable")

    with open(os.path.join(OUT, "Pacific_Pack_Global_email.eml"), "wb") as handle:
        handle.write(bytes(message))


# ---------------------------------------------------------------------------
# the buyer's own item list, written out of the catalogue above
#
# The standing items existed in two places -- here, and again as `lines` in
# examples/corrugated_packaging.json -- typed out twice and identical only for
# as long as nobody edited one of them. That is precisely the drift the example
# file's own docstring warns about: the seeded request and the fabricated
# replies quoting different catalogues, which opens the demo on thirty rows of
# "not quoted". So the JSON's item rows are now written from CATALOGUE, and
# everything else in that file -- the title, terms, questions and the optional
# extras a buyer might add in a given year -- is left exactly as it is.
# ---------------------------------------------------------------------------

UNIT_FOR_FAMILY = {"sheet": "per sheet", "roll": "per roll",
                   "consumable": "per roll"}


def catalogue_rows() -> list[dict]:
    """CATALOGUE as the example request carries it."""
    return [{"sku": item.sku,
             "description": item.description,
             "quantity": item.quantity,
             "unit": UNIT_FOR_FAMILY.get(item.family, "per box"),
             "moq": item.moq}
            for item in CATALOGUE]


def write_example_request() -> None:
    import json

    path = os.path.join(ROOT, "examples", "corrugated_packaging.json")
    with open(path, encoding="utf-8") as handle:
        data = json.load(handle)
    data["lines"] = catalogue_rows()
    with open(path, "w", encoding="utf-8") as handle:
        json.dump(data, handle, indent=2, ensure_ascii=False)
        handle.write("\n")


# ---------------------------------------------------------------------------
# last year's purchase order -- an input, not a supplier response
#
# Lives in examples/ rather than sample_data/ because it is something the BUYER
# already has. It exists so the drafting page's "attach a list you already
# have" path can be demonstrated on a real document: fifteen lines, last
# year's prices, last year's date, and a supplier's name on it.
# ---------------------------------------------------------------------------

def build_purchase_order() -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)

    examples_dir = os.path.join(ROOT, "examples")
    os.makedirs(examples_dir, exist_ok=True)
    path = os.path.join(examples_dir, "PO_2025_Corrugated_Packaging.pdf")

    base = getSampleStyleSheet()
    ink = colors.HexColor("#1A1A1A")
    body = ParagraphStyle("b", parent=base["Normal"], fontSize=8.6, leading=12)
    small = ParagraphStyle("s", parent=base["Normal"], fontSize=7.4, leading=10,
                           textColor=colors.HexColor("#555555"))
    cell = ParagraphStyle("c", parent=base["Normal"], fontSize=7.6, leading=10)
    cellb = ParagraphStyle("cb", parent=cell, fontName="Helvetica-Bold")
    head = ParagraphStyle("h", parent=base["Normal"], fontName="Helvetica-Bold",
                          fontSize=13, leading=16, textColor=ink)

    doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=18 * mm,
                            rightMargin=18 * mm, topMargin=16 * mm,
                            bottomMargin=16 * mm, title="Purchase Order")

    story = [
        Paragraph(brand.COMPANY_LEGAL, head),
        Paragraph(f"{brand.ADDRESS} · GSTIN 29AAJCW8821K1Z4", small),
        Spacer(1, 10),
        Paragraph("PURCHASE ORDER", ParagraphStyle(
            "po", parent=head, fontSize=11, textColor=colors.HexColor("#22417C"))),
        Spacer(1, 4),
    ]

    facts = [
        ["PO number", "WCB/PO/2025-26/0417"],
        ["Date", PO_DATE],
        ["Supplier", "Shakti Packaging Industries Pvt Ltd, Bengaluru"],
        ["Buyer contact", f"{brand.CONTACT} · {brand.CONTACT_EMAIL}"],
        ["Deliver to", brand.ADDRESS],
        ["Payment terms", "45 days from date of invoice"],
        ["Contract period", "FY 2025-26"],
    ]
    facts_table = Table([[Paragraph(f"<b>{k}</b>", cell), Paragraph(v, cell)]
                         for k, v in facts], colWidths=[30 * mm, 144 * mm])
    facts_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#DDDDDD")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(facts_table)
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "The following items were ordered against the annual rate contract for "
        "corrugated packaging. Rates below are the FY 2025-26 contracted rates.",
        body))
    story.append(Spacer(1, 6))

    rows = [[Paragraph(t, cellb) for t in
             ("Sl", "Item code", "Description", "UOM", "Quantity",
              "Rate (INR)", "Value (INR)")]]
    total = 0.0
    # Fifteen of the thirty, so the drafted request from this PO is visibly a
    # subset -- a buyer starting from last year's order does not get every line
    # back, and should notice which ones are missing.
    for serial, item in enumerate(CATALOGUE[:15], start=1):
        uom = {"sheet": "Sheet", "roll": "Roll", "consumable": "Roll"}.get(
            item.family, "Nos")
        rate = round(PRICES["shakti"][item.sku] * 0.965, 2)   # last year's rate
        value = round(rate * item.quantity, 2)
        total += value
        rows.append([
            Paragraph(str(serial), cell),
            Paragraph(item.sku, cell),
            Paragraph(item.description, cell),
            Paragraph(uom, cell),
            Paragraph(f"{item.quantity:,}", cell),
            Paragraph(f"{rate:,.2f}", cell),
            Paragraph(f"{value:,.2f}", cell),
        ])
    rows.append([Paragraph("", cell), Paragraph("", cell),
                 Paragraph("<b>Total (excluding GST)</b>", cell),
                 Paragraph("", cell), Paragraph("", cell), Paragraph("", cell),
                 Paragraph(f"<b>{total:,.2f}</b>", cell)])

    table = Table(rows, colWidths=[9 * mm, 20 * mm, 66 * mm, 13 * mm, 20 * mm,
                                   22 * mm, 24 * mm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (4, 1), (-1, -1), "RIGHT"),
        ("LINEABOVE", (0, -1), (-1, -1), 0.9, ink),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(table)

    story.append(Spacer(1, 10))
    story.append(Paragraph("Terms", ParagraphStyle(
        "t2", parent=head, fontSize=9.5, leading=13)))
    for line in (
        "1. Rates are firm for the contract period and inclusive of delivery to "
        "the address above.",
        "2. GST extra as applicable. HSN code to be stated on every invoice.",
        "3. Goods to be supplied against release schedules issued monthly.",
        "4. Rejections to be replaced within 7 working days at supplier's cost.",
    ):
        story.append(Paragraph(line, body))

    story.append(Spacer(1, 14))
    story.append(Paragraph(f"For {brand.COMPANY_LEGAL}", body))
    story.append(Spacer(1, 16))
    story.append(Paragraph(brand.CONTACT, body))

    doc.build(story)
    print(f"examples/\n  {os.path.basename(path)}"
          f"{'.' * max(1, 46 - len(os.path.basename(path)))}"
          f"{os.path.getsize(path) / 1024:6.1f} KB")

# ---------------------------------------------------------------------------
# supporting documents the quotations refer to
# ---------------------------------------------------------------------------

def build_attachments() -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)

    os.makedirs(ATTACH, exist_ok=True)
    styles = getSampleStyleSheet()
    small = ParagraphStyle("s", parent=styles["Normal"], fontSize=9, leading=13)
    title = ParagraphStyle("t", parent=styles["Normal"], fontSize=15, leading=19,
                           spaceAfter=8)

    def make(filename: str, blocks: list) -> None:
        SimpleDocTemplate(os.path.join(ATTACH, filename), pagesize=A4,
                          leftMargin=20 * mm, rightMargin=20 * mm,
                          topMargin=20 * mm, bottomMargin=18 * mm).build(blocks)

    make("Shakti_ISO9001_certificate.pdf", [
        Paragraph("<b>CERTIFICATE OF REGISTRATION</b>", title),
        Paragraph("ISO 9001:2015 — Quality Management Systems", small),
        Spacer(1, 6 * mm),
        Paragraph("This is to certify that<br/><br/><b>SHAKTI PACKAGING INDUSTRIES "
                  "PVT LTD</b><br/>Plot 44, Phase II, Peenya Industrial Area, "
                  "Bengaluru 560058, India", small),
        Spacer(1, 4 * mm),
        Paragraph("operates a Quality Management System which complies with the "
                  "requirements of ISO 9001:2015 for the following scope:<br/><br/>"
                  "<i>Manufacture and supply of corrugated fibreboard boxes, sheets "
                  "and allied packaging products.</i>", small),
        Spacer(1, 6 * mm),
        Paragraph("Certificate No.: IN-QMS-2023-77410<br/>"
                  "Original issue: 14 November 2023<br/>"
                  "Current issue: 14 November 2025<br/>"
                  "Valid until: 13 November 2027<br/>"
                  "Issued by: Trident Certification Services, Chennai", small),
    ])

    rows = [["Test", "Method", "Specified", "Result", "Verdict"],
            ["Bursting strength", "IS 7063", "≥ 16 kg/cm²", "17.4 kg/cm²", "Pass"],
            ["Edge crush (ECT)", "IS 14871", "≥ 5.2 kN/m", "5.8 kN/m", "Pass"],
            ["Grammage, outer liner", "IS 1060", "180 ± 8 GSM", "183 GSM", "Pass"],
            ["Moisture content", "IS 1060", "≤ 9%", "7.2%", "Pass"],
            ["Drop test, 10 drops", "IS 7028", "No rupture", "No rupture", "Pass"],
            ["Compression, stacked", "IS 7028", "≥ 250 kgf", "268 kgf", "Pass"]]
    table = Table(rows, colWidths=[45 * mm, 28 * mm, 32 * mm, 30 * mm, 22 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
    ]))
    make("SBC_burst_strength_test_report.pdf", [
        Paragraph("<b>THIRD PARTY TEST REPORT</b>", title),
        Paragraph("SGS India Pvt Ltd, Bengaluru Laboratory<br/>"
                  "Report No. SGS/BLR/2026/08/2214 &nbsp;|&nbsp; Date: 02 August 2026",
                  small),
        Spacer(1, 5 * mm),
        Paragraph("Client: Sri Balaji Corrugators Pvt Ltd<br/>"
                  "Sample: 5-ply corrugated box, 500x400x350 mm, BF-18 grade<br/>"
                  "Sample received: 28 July 2026", small),
        Spacer(1, 5 * mm), table, Spacer(1, 5 * mm),
        Paragraph("Results relate only to the sample tested. This report shall not be "
                  "reproduced except in full.", small),
    ])

    make("Meridian_AnnexureA_certificates.pdf", [
        Paragraph("<b>ANNEXURE A — CERTIFICATIONS AND TEST REPORTS</b>", title),
        Paragraph("Meridian Packaging LLP", small),
        Spacer(1, 5 * mm),
        Paragraph("1. ISO 9001:2015, certificate MPL/QMS/2025/119, valid to "
                  "31 March 2028, issued by Bureau Veritas India.", small),
        Spacer(1, 3 * mm),
        Paragraph("2. Third-party burst strength report, Intertek Bengaluru, "
                  "ref ITK/2026/06/8841 dated 19 June 2026 — 5-ply BF-16, "
                  "result 16.9 kg/cm², pass.", small),
        Spacer(1, 3 * mm),
        Paragraph("3. Third-party compression report, Intertek Bengaluru, "
                  "ref ITK/2026/06/8842 dated 19 June 2026 — stacked compression "
                  "254 kgf against 250 kgf specified, pass.", small),
        Spacer(1, 3 * mm),
        Paragraph("4. FSC Chain of Custody, licence FSC-C148820, covering recycled "
                  "kraft liner, valid to 08 February 2027.", small),
        Spacer(1, 6 * mm),
        Paragraph("<i>Note: our on-time delivery performance for FY 2025-26 was "
                  "91.8%. The corrective action plan agreed with our largest customer "
                  "is available on request.</i>", small),
    ])


# ---------------------------------------------------------------------------

def main() -> None:
    os.makedirs(OUT, exist_ok=True)
    for name in os.listdir(OUT):
        path = os.path.join(OUT, name)
        if os.path.isfile(path):
            os.remove(path)

    build_shakti()
    build_balaji()
    build_meridian()
    build_northstar()
    build_pacific()
    build_attachments()
    build_purchase_order()
    write_example_request()

    print("sample_data/")
    for name in sorted(os.listdir(OUT)):
        path = os.path.join(OUT, name)
        if os.path.isfile(path):
            print(f"  {name:44} {os.path.getsize(path) / 1024:8.1f} KB")
    print("sample_data/attachments/")
    for name in sorted(os.listdir(ATTACH)):
        print(f"  {name:44} "
              f"{os.path.getsize(os.path.join(ATTACH, name)) / 1024:8.1f} KB")


if __name__ == "__main__":
    main()
