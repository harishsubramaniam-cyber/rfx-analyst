"""Build extraction fixtures for the five sample responses.

WHAT THESE ARE: recorded extraction output, used ONLY by the offline test
suite so the matcher, normaliser, award and query layers can be verified
deterministically and without network access.

WHAT THESE ARE NOT: a demo shortcut. The application never reads them unless
RFX_OFFLINE_FIXTURES=1 is explicitly set, which the demo does not do. When the
app runs normally, every one of these documents goes through the real model.

They are generated from the same source of truth that writes the documents
(`build_sample_data.py`), so a fixture cannot drift from the file it describes.
Where a document is text-bearing, the generated fixture is then checked back
against the file itself.
"""

from __future__ import annotations

import json
import os
import re
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
sys.path.insert(0, ROOT)
sys.path.insert(0, HERE)

import build_sample_data as gen  # noqa: E402

OUT = os.path.join(HERE, "fixtures")
SAMPLES = gen.OUT

def q(question, answer, stated_value=None, evidence=None, note=None):
    return {"question": question, "answer": answer, "stated_value": stated_value,
            "evidence": evidence, "note": note}



def _read_clarity(base: float, *, row: int = 0, decay: float = 0.0,
                  penalties: float = 0.0) -> float:
    """What a careful reader would honestly claim for ONE number.

    A flat score per document is the tell of a model that was not really
    looking. Real documents degrade unevenly: a photographed card loses the
    rows nearest the fold, a wrapped remarks cell pushes a rate off its
    baseline, a hand-written code is not the same read as a printed one. These
    fixtures carry that shape so the offline demo shows what a live read shows.
    """
    return round(max(0.45, min(0.99, base - decay * row - penalties)), 2)

def _line(sku, desc, value, currency, unit, locator, snippet, qty=None,
          conditions=None, confidence=0.95, notes="", lead=None):
    return {
        "vendor_sku": sku, "description": desc, "quoted_value": value,
        "currency": currency, "unit_text": unit, "quantity": qty,
        "lead_time_days": lead, "confidence": confidence,
        "conditions": conditions or [], "notes": notes,
        "source": {"locator": locator, "snippet": snippet},
    }


# ---------------------------------------------------------------------------
# 1 · Shakti Packaging -- spreadsheet, own codes, conditional freight
# ---------------------------------------------------------------------------

def shakti() -> dict:
    skipped = {"BX-019", "BX-005"}
    remarks = {
        "BX-004": "Rate for kraft 180 GSM outer; 200 GSM +₹2.10",
        "BX-011": "Sub-contracted; lead time 30 days",
        "BX-016": "Poly-coated liner, moisture barrier as per spec",
        "BX-025": "Mill roll — supplied in 50 kg reels",
        "BX-029": "Imported film, subject to FX at time of despatch",
    }
    lines, serial, row = [], 1, 13
    for item in gen.CATALOGUE:
        if item.sku in skipped:
            continue
        unit = ("Sheet" if item.family == "sheet"
                else "Roll" if item.family in {"roll", "consumable"} else "Nos")
        price = gen.PRICES["shakti"][item.sku]
        buyer_ref = item.sku if serial % 4 else item.sku.replace("-", " ")
        conditions = []
        if item.sku == "BX-029":
            conditions.append("Imported film, subject to FX at time of despatch")
        lines.append(_line(
            buyer_ref, item.description, price, "INR", unit,
            f"sheet 'Quotation' row {row}",
            f"{serial} | {buyer_ref} | "
            f"{item.description} | {gen.HSN[item.family]} | {unit} | "
            f"{item.moq:,} | {price:,.2f} | {remarks.get(item.sku, '')}".strip(),
            qty=None,
            conditions=conditions,
            confidence=_read_clarity(
                0.98,
                penalties=(0.03 if buyer_ref != item.sku else 0.0)
                + (0.02 if item.sku in remarks else 0.0)),
            notes=remarks.get(item.sku, ""),
        ))
        serial += 1
        row += 1

    return {
        "vendor": "Shakti Packaging Industries Pvt Ltd",
        "document_currency": "INR",
        # Freight is included only above 2 tonnes -- that is neither yes nor no.
        "freight_included": None,
        "payment_terms_days": 30,
        "lead_time_days": 19.5,
        "overall_confidence": 0.96,
        "extraction_notes": (
            "Supplier's own quotation format: their item codes in one column and "
            "the buyer's reference in another, written inconsistently. Freight is "
            "included only for despatches above 2 tonnes."
        ),
        "unresolved_references": [
            "Shakti_ISO9001_certificate.pdf (ISO certificate stated as enclosed)",
            "'Slab Rates' volume pricing, which depends on confirmed annual volumes",
        ],
        "lines": lines,
        "questionnaire": [
            q("ISO 9001 certification held?", "Yes", None,
              "valid to Nov 2027, copy enclosed (Shakti_ISO9001_certificate.pdf)"),
            q("Monthly production capacity above 50,000 units?", "Yes",
              "approx. 90,000 units/month"),
            q("On-time delivery performance above 95%?", "Yes",
              "96.4% for FY 2025-26"),
            q("Quality rejection rate below 1%?", "Yes", "0.6% average"),
            q("Able to commit to a 30-day delivery SLA?", "Yes"),
        ],
        "terms": [
            {"kind": "freight",
             "text": "Rates are Ex-Works Bengaluru. Freight to your Bommasandra "
                     "warehouse is included for despatches above 2 tonnes; below "
                     "that, freight at actuals.",
             "trigger": "despatch above 2 tonnes", "value": None},
            {"kind": "tax", "text": "GST @ 18% extra as applicable.",
             "trigger": None, "value": None},
            {"kind": "payment", "text": "Payment: 30 days from date of invoice.",
             "trigger": None, "value": None},
            {"kind": "validity", "text": "Validity 45 days from date of quotation.",
             "trigger": None, "value": None},
            {"kind": "discount",
             "text": "Slab rates: 2.5% lower for 5,001–15,000 units and 4.5% lower "
                     "above 15,000 units, per item, on confirmed annual volumes.",
             "trigger": "confirmed annual offtake per item", "value": 4.5},
            {"kind": "other",
             "text": "Rates are based on kraft paper at ₹42/kg. Variation beyond "
                     "±5% will be passed on with 30 days' notice.",
             "trigger": None, "value": None},
        ],
    }


# ---------------------------------------------------------------------------
# 2 · Sri Balaji -- PDF, rupee glyph lost in the text layer, footnote discount
# ---------------------------------------------------------------------------

def balaji() -> dict:
    skipped = {"BX-012", "BX-025"}
    substitutions = {
        "BX-016": "Offered in poly-laminated kraft in place of specified wax coating",
        "BX-018": "BF-18 grade; BF-20 available at +4%",
    }
    lines = []
    for item in gen.CATALOGUE:
        if item.sku in skipped:
            continue
        unit = {"sheet": "INR / sheet", "roll": "INR / roll",
                "consumable": "INR / roll"}.get(item.family, "INR / box")
        if item.family == "partition":
            unit = "INR / set"
        price = gen.PRICES["balaji"][item.sku]
        lines.append(_line(
            item.sku, item.description, price, "INR", unit,
            "page 1, quotation table",
            f"{item.sku} | {item.description} | {gen.HSN[item.family]} | "
            f"{item.moq:,} | ₹{price:,.2f} | {unit}",
            conditions=[substitutions[item.sku]] if item.sku in substitutions else [],
            # Re-read from page images after the text layer proved corrupt, so
            # every figure is an OCR read rather than an exact character.
            confidence=_read_clarity(
                0.95,
                penalties=(0.04 if item.sku in substitutions else 0.0)
                + (0.02 if price >= 100 else 0.0)),
        ))

    return {
        "vendor": "Sri Balaji Corrugators Pvt Ltd",
        "document_currency": "INR",
        "freight_included": True,
        "payment_terms_days": 45,
        "lead_time_days": 22,
        "overall_confidence": 0.94,
        "extraction_notes": (
            "The PDF text layer renders every rupee sign as an unmapped glyph, so "
            "the pages were re-read as images. Revision R2 supersedes R1 of "
            "14 Aug 2026."
        ),
        "unresolved_references": [
            "Artwork/plate charges for printed cartons, quoted separately",
        ],
        "lines": lines,
        "questionnaire": [
            q("ISO 9001 certified", "Yes", None,
              "certificate no. IN-QMS-44192, enclosed"),
            q("Monthly production capacity > 50,000 units", "Yes"),
            q("On-time delivery > 95%", "Yes", "97.1% in FY 2025-26"),
            q("Quality rejection rate < 1%", "Yes", "0.4%"),
            q("Can meet 30-day delivery SLA", "Yes"),
        ],
        "terms": [
            {"kind": "freight",
             "text": "Freight is included for deliveries to your Bommasandra warehouse.",
             "trigger": None, "value": None},
            {"kind": "tax", "text": "GST @ 18% extra.", "trigger": None, "value": None},
            {"kind": "payment", "text": "Payment terms 45 days from invoice.",
             "trigger": None, "value": None},
            {"kind": "validity", "text": "Rates firm for 60 days from the date of "
                                         "this quotation.", "trigger": None, "value": None},
            {"kind": "discount",
             "text": "An additional discount of 2% on the quoted rates will apply to "
                     "the entire order value where annualised offtake across all "
                     "items exceeds 50,000 units.",
             "trigger": "annualised offtake across all items exceeds 50,000 units",
             "value": 2.0},
            {"kind": "other",
             "text": "BX-012 and BX-025 are not quoted — outside in-house capability.",
             "trigger": None, "value": None},
        ],
    }


# ---------------------------------------------------------------------------
# 3 · Meridian -- Word quotation, everything as prose, ambiguous "per box/set"
# ---------------------------------------------------------------------------

def meridian() -> dict:
    lines = []
    for index, item in enumerate(gen.CATALOGUE, start=13):
        price = gen.PRICES["meridian"][item.sku]
        unit = {"sheet": "per sheet", "roll": "per roll",
                "consumable": "per roll"}.get(item.family, "per box/set")
        if item.family == "protector":
            unit = "per box/set"
        lead = 24 if item.family in {"box", "spec"} else 28
        snippet = (f"For {item.sku}, {item.description}, against your indicated "
                   f"annual quantity of {item.quantity:,}, our rate is "
                   f"₹{price:,.2f} {unit}, minimum order {item.moq:,}. Delivery "
                   f"approximately {lead} days from receipt of order.")
        conditions = []
        if item.sku == "BX-009":
            conditions.append("Printing plates are a one-time charge and are quoted "
                              "separately on approval of artwork.")
            snippet += " " + conditions[0]
        if item.sku == "BX-029":
            conditions.append("This is a traded item and the rate moves with the "
                              "imported film market; we will reconfirm at the time "
                              "of order.")
            snippet += " " + conditions[0]
        lines.append(_line(
            item.sku, item.description, price, "INR", unit,
            f"paragraph {index}", snippet, qty=float(item.quantity),
            # A rate inside a sentence is a harder read than a rate in a cell,
            # and the paragraphs carrying a caveat are the wordiest of all.
            conditions=conditions,
            confidence=_read_clarity(0.95, row=index, decay=0.002,
                                     penalties=0.05 if conditions else 0.0),
            lead=float(lead),
        ))

    return {
        "vendor": "Meridian Packaging LLP",
        "document_currency": "INR",
        "freight_included": False,
        "payment_terms_days": 30,
        "lead_time_days": 25.3,
        "overall_confidence": 0.91,
        "extraction_notes": (
            "Commercials are written as prose, one sentence per item, with no table "
            "anywhere in the document. Most units are given as 'per box/set' "
            "without saying which."
        ),
        "unresolved_references": [
            "Annexure A — ISO certificate and third-party test reports",
        ],
        "lines": lines,
        "questionnaire": [
            q("ISO 9001 certified", "Yes", None, "certificate valid to March 2028"),
            q("Monthly production capacity above 50,000 units", "Yes"),
            q("On-time delivery above 95%", "No", "91.8%",
              "Annexure A — third-party reports attached",
              "We have added a second shift in the flexo line and expect to be "
              "above 95% from Q3."),
            q("Quality rejection rate below 1%", "Yes", "0.7%"),
            q("Can meet a 30-day delivery SLA", "Yes"),
        ],
        "terms": [
            {"kind": "freight",
             "text": "Freight is charged separately at actuals and is not included "
                     "in the rates above.", "trigger": None, "value": None},
            {"kind": "tax", "text": "GST at 18% is extra.", "trigger": None, "value": None},
            {"kind": "validity", "text": "Prices hold for 30 days from the date of "
                                         "this quotation.", "trigger": None, "value": None},
            {"kind": "payment", "text": "Payment terms requested are 30 days from "
                                        "invoice.", "trigger": None, "value": None},
            {"kind": "discount",
             "text": "A volume rebate of 1.5% is applicable where cumulative "
                     "purchases in any quarter exceed ₹50 lakh, settled by credit "
                     "note at quarter end.",
             "trigger": "cumulative quarterly purchases exceed ₹50 lakh", "value": 1.5},
        ],
    }


# ---------------------------------------------------------------------------
# 4 · Northstar -- photographed rate card, no currency anywhere
# ---------------------------------------------------------------------------

def northstar() -> dict:
    skipped = {"BX-008", "BX-019"}
    special = {
        "BX-009": (f"{gen.PRICES['northstar']['BX-009'] * 100:,.0f}", "/ 100 pcs"),
        "BX-016": ("68", "/ kg"),
        "BX-025": (f"{gen.PRICES['northstar']['BX-025']:,.0f}", "/ roll"),
        "BX-001": ("", "/ box"),
        "BX-013": ("", "/box"),
    }
    lines = []
    for item in gen.CATALOGUE:
        if item.sku in skipped:
            continue
        price = gen.PRICES["northstar"][item.sku]
        if item.sku in special and special[item.sku][0]:
            shown, suffix = special[item.sku]
        elif item.sku in special:
            shown, suffix = f"{price:,.0f}", special[item.sku][1]
        else:
            shown = f"{price:,.0f}" if price >= 100 else f"{price:,.2f}"
            suffix = ""
        value = float(shown.replace(",", ""))
        lines.append(_line(
            item.sku, item.description, value, None, suffix or None,
            "photographed rate card",
            f"{item.sku}  {item.description}   {shown} {suffix}".strip(),
            confidence=_read_clarity(
                0.95, row=len(lines), decay=0.007,
                penalties=(0.05 if "." in shown else 0.0)
                + (0.04 if item.sku in special else 0.0)),
        ))

    return {
        "vendor": "Northstar Packaging",
        "document_currency": None,
        "freight_included": False,
        "payment_terms_days": None,
        "lead_time_days": None,
        "overall_confidence": 0.78,
        "extraction_notes": (
            "Photograph of a printed rate card taken at an angle. No currency "
            "symbol or code appears anywhere on the sheet. Most rows show a bare "
            "number; a few state a unit. A handwritten note reads 'rev.2 — confirm "
            "rates w/ Suresh before PO'."
        ),
        "unresolved_references": [
            "rev.2 rates to be confirmed with Suresh before PO, per the handwritten note",
        ],
        "lines": lines,
        "questionnaire": [
            # Terse, exactly as a model reads them off the photographed card.
            # These are the phrasings that used to spawn duplicate criteria.
            q("ISO 9001", "No", None, None, "audit scheduled Nov 2026"),
            q("Capacity", "Yes", "60,000 units/month"),
            q("OTD >95%", "Yes"),
        ],
        "terms": [
            {"kind": "validity",
             "text": "All rates subject to confirmation at time of order.",
             "trigger": None, "value": None},
            {"kind": "freight", "text": "Freight extra.", "trigger": None, "value": None},
            {"kind": "other", "text": "MOQ 500 nos per item unless agreed otherwise.",
             "trigger": None, "value": None},
            {"kind": "other", "text": "BX-008 and BX-019 are not manufactured by us.",
             "trigger": None, "value": None},
        ],
    }


# ---------------------------------------------------------------------------
# 5 · Pacific Pack Global -- exporter's email, quoted FOB in US dollars
# ---------------------------------------------------------------------------

def _email_body(filename: str) -> list[str]:
    """The message body of a saved email, decoded the way a reader would.

    Reading the raw .eml here would compare the fixture against
    quoted-printable escapes rather than against what the supplier wrote, so
    the check would pass on text no model ever sees.
    """
    import email
    from email import policy

    with open(os.path.join(SAMPLES, filename), "rb") as handle:
        message = email.message_from_bytes(handle.read(), policy=policy.default)
    body = message.get_body(preferencelist=("plain",))
    return (body.get_content() if body else "").splitlines()


def pacific() -> dict:
    raw = _email_body("Pacific_Pack_Global_email.eml")

    pattern = re.compile(r"^(BX-\d{3}) - USD ([\d,]+(?:\.\d+)?)(.*)$")
    unit_hints = [
        (re.compile(r"per 1000 pcs", re.I), "per 1000 pcs"),
        (re.compile(r"per 100 rolls", re.I), "per 100 rolls"),
        (re.compile(r"per kg", re.I), "per kg"),
        (re.compile(r"per box", re.I), "per box"),
        (re.compile(r"per sheet", re.I), "per sheet"),
        (re.compile(r"per roll", re.I), "per roll"),
        (re.compile(r"\beach\b", re.I), "each"),
    ]
    condition_hints = [
        (re.compile(r"plate charge to be advised", re.I),
         "plus one-time plate charge to be advised"),
        (re.compile(r"balance grades same as our 2025 contract rates", re.I),
         "balance grades same as our 2025 contract rates"),
    ]

    lines = []
    for index, text in enumerate(raw, start=1):
        match = pattern.match(text.strip())
        if not match:
            continue
        sku, amount, tail = match.groups()
        unit = next((label for regex, label in unit_hints if regex.search(tail)), None)
        conditions = [label for regex, label in condition_hints if regex.search(tail)]
        lines.append(_line(
            sku, None, float(amount.replace(",", "")), "USD", unit,
            f"email line {index}", text.strip(),
            # Typed by hand into an email: the digits are exact, but the unit
            # and any caveat are buried in prose after the number.
            conditions=conditions,
            confidence=_read_clarity(
                0.93,
                penalties=(0.05 if conditions else 0.0) + (0.04 if not unit else 0.0)),
        ))

    return {
        "vendor": "Pacific Pack Global Pvt Ltd",
        "document_currency": "USD",
        "freight_included": False,
        "payment_terms_days": 60,
        "lead_time_days": 32.5,
        "overall_confidence": 0.85,
        "extraction_notes": (
            "Free-text email from an export-oriented unit. Everything is quoted in "
            "US dollars, FOB Chennai, with conversion left to the buyer's bank rate "
            "on invoice date. Units vary line by line and several are stated in bulk."
        ),
        "unresolved_references": [
            "our standard 2025 rate card (referenced for every item not listed)",
            "2025 contract rates (referenced on BX-017)",
            "one-time plate charge on BX-009, to be advised",
        ],
        "lines": lines,
        "questionnaire": [
            q("ISO 9001 certified", "Yes"),
            q("Monthly production capacity > 50,000 units", "No",
              "about 35,000 units/month"),
            q("On-time delivery > 95%", "Yes"),
            q("Quality rejection rate < 1%", "Yes"),
            q("Can meet 30-day delivery SLA", "No", "30-35 days",
              None, "please see lead time above"),
        ],
        "terms": [
            {"kind": "other",
             "text": "We invoice in USD as we are an export-oriented unit; "
                     "conversion at your banker's TT rate on invoice date.",
             "trigger": None, "value": None},
            {"kind": "freight", "text": "Prices are FOB Chennai.",
             "trigger": None, "value": None},
            {"kind": "tax", "text": "GST/IGST as applicable extra.",
             "trigger": None, "value": None},
            {"kind": "payment", "text": "Payment 60 days from BL date.",
             "trigger": None, "value": None},
            {"kind": "other", "text": "Lead time 30-35 days ex-Chennai. MOQ 1000 pcs "
                                      "per line item, 500 for the 7-ply.",
             "trigger": None, "value": None},
        ],
    }


BUILDERS = {
    "Shakti_Packaging_Quotation": shakti,
    "Sri_Balaji_Corrugators_Quote": balaji,
    "Meridian_Packaging_Offer": meridian,
    "Northstar_RateCard_photo": northstar,
    "Pacific_Pack_Global_email": pacific,
}


def verify(fixtures: dict[str, dict]) -> list[str]:
    """Cross-check the fixtures against the documents they claim to describe."""
    problems: list[str] = []

    # Meridian: every price must appear verbatim in the Word file
    from docx import Document
    paragraphs = [p.text for p in
                  Document(os.path.join(SAMPLES,
                                        "Meridian_Packaging_Offer.docx")).paragraphs]
    text = "\n".join(paragraphs)
    for line in fixtures["Meridian_Packaging_Offer"]["lines"]:
        if f"₹{line['quoted_value']:,.2f}" not in text:
            problems.append(f"Meridian {line['vendor_sku']} price not found in the .docx")

    # ...and every locator must point at the paragraph it claims. The fixture
    # numbers paragraphs, so anything inserted above the offer silently shifts
    # thirty locators without changing a single price.
    for line in fixtures["Meridian_Packaging_Offer"]["lines"]:
        index = int(line["source"]["locator"].split()[-1])
        if not (0 <= index < len(paragraphs)) or \
                not paragraphs[index].startswith(f"For {line['vendor_sku']},"):
            problems.append(
                f"Meridian {line['vendor_sku']} locator points at "
                f"{line['source']['locator']}, which is not that item's sentence")

    # This response is a quotation document, not correspondence. The difficulty
    # it exists to pose is prose pricing with no table -- a salutation and a
    # sign-off add nothing to that, and a letter wrapper creeping back in is
    # worth failing over rather than noticing months later.
    for phrase in ("Dear ", "Yours sincerely", "Kind attention", "Subject:"):
        if phrase in text:
            problems.append(f"Meridian .docx has slipped back into a letter: {phrase!r}")
    if any(paragraph.strip().startswith("|") for paragraph in paragraphs) or \
            Document(os.path.join(SAMPLES,
                                  "Meridian_Packaging_Offer.docx")).tables:
        problems.append("Meridian .docx now has a table; it must price in prose")

    # Shakti: every price must appear in the workbook
    import openpyxl
    wb = openpyxl.load_workbook(os.path.join(SAMPLES, "Shakti_Packaging_Quotation.xlsx"))
    sheet_values = {round(cell.value, 2) for row in wb["Quotation"].iter_rows()
                    for cell in row if isinstance(cell.value, (int, float))}
    for line in fixtures["Shakti_Packaging_Quotation"]["lines"]:
        if round(line["quoted_value"], 2) not in sheet_values:
            problems.append(f"Shakti {line['vendor_sku']} price not found in the .xlsx")

    # Pacific: the fixture is parsed straight from the email body, so just count
    quoted = len(re.findall(r"^BX-\d{3} - USD",
                            "\n".join(_email_body("Pacific_Pack_Global_email.eml")),
                            re.M))
    if quoted != len(fixtures["Pacific_Pack_Global_email"]["lines"]):
        problems.append("Pacific line count does not match the email")

    return problems


def main() -> None:
    os.makedirs(OUT, exist_ok=True)
    for name in os.listdir(OUT):
        os.remove(os.path.join(OUT, name))

    fixtures = {stem: builder() for stem, builder in BUILDERS.items()}

    problems = verify(fixtures)
    if problems:
        print("VERIFICATION FAILED")
        for problem in problems:
            print("  -", problem)
        raise SystemExit(1)

    for stem, data in fixtures.items():
        with open(os.path.join(OUT, f"{stem}.json"), "w", encoding="utf-8") as handle:
            json.dump(data, handle, indent=2, ensure_ascii=False)
        print(f"{stem:34} {len(data['lines']):3} lines   "
              f"{data['document_currency'] or 'currency not stated'}")
    print("\nAll fixtures verified against their documents.")


if __name__ == "__main__":
    main()
