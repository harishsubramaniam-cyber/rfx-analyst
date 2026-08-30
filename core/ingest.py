"""Turn any vendor file into something the model can read.

Design notes worth keeping in mind:

* Text is emitted WITH locators (sheet + row, paragraph index, page number) so
  the extractor can cite a real position, not just quote a fragment.

* PDF text extraction is checked before it is trusted. Embedded-font quirks
  silently destroy currency symbols -- the sample ABC Corrugators quote
  extracts every price as "■40.74" through pypdf, losing the rupee sign
  entirely. When that is detected we re-render the pages as images and let the
  vision model read the page as a human would.

* Photographs are a first-class input, not an afterthought. A phone snap of a
  printed rate card is one of the five formats real buyers actually receive.
"""

from __future__ import annotations

import io
import os
import re
from typing import Optional

import pandas as pd
from docx import Document
from pypdf import PdfReader

from .models import DocumentPayload

IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp", "heic", "bmp", "tif", "tiff"}
TEXT_EXTENSIONS = {"txt", "csv", "md"}

# How many messages deep a forwarded chain may be unpacked. See read_eml.
MAX_EMAIL_DEPTH = 3

# What the vision model will actually accept. Anything else is converted to
# PNG before it is sent.
VISION_MIMES = {"image/jpeg", "image/png", "image/webp"}

# Characters that mean the font mapping failed rather than the page being empty.
_CORRUPTION_MARKERS = ("�", "■", "□", "�")


def extension_of(name: str) -> str:
    return name.lower().rsplit(".", 1)[-1] if "." in name else ""


# ---------------------------------------------------------------------------
# spreadsheets
# ---------------------------------------------------------------------------

def read_spreadsheet(handle, filename: str) -> DocumentPayload:
    sheets = pd.read_excel(handle, sheet_name=None, header=None)
    chunks: list[str] = []
    for sheet_name, frame in sheets.items():
        chunks.append(f"--- SHEET: {sheet_name} ---")
        for row_index, row in frame.iterrows():
            cells = [("" if pd.isna(v) else str(v)).strip() for v in row.tolist()]
            if not any(cells):
                continue
            chunks.append(f"[row {row_index + 1}] " + " | ".join(cells))
    return DocumentPayload(
        file=filename,
        text="\n".join(chunks),
        locator_hint="sheet '<name>' row <n>",
        reader="pandas/openpyxl",
    )


def _decode(raw) -> str:
    """Bytes to text, without silently eating the currency symbol.

    `decode("utf-8", errors="ignore")` looks tolerant and is in fact the exact
    failure this module exists to prevent: a CSV saved by Excel on a Windows
    machine is cp1252, where ₹ and £ are single bytes that are not valid UTF-8,
    so "ignore" deletes them. "£4.20 per kg" arrives as "4.20 per kg", the
    document then states no currency anywhere, and every price on it is flagged
    as assumed. Try the encodings a supplier's file is actually saved in, in
    order, and only fall back to discarding bytes if none of them decode.
    """
    if not isinstance(raw, bytes):
        return raw or ""
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def read_csv(handle, filename: str) -> DocumentPayload:
    raw = _decode(handle.read())
    lines = [f"[row {i + 1}] {line}" for i, line in enumerate(raw.splitlines()) if line.strip()]
    return DocumentPayload(file=filename, text="\n".join(lines),
                           locator_hint="row <n>", reader="text")


# ---------------------------------------------------------------------------
# word documents
# ---------------------------------------------------------------------------

def read_docx(handle, filename: str) -> DocumentPayload:
    document = Document(handle)
    chunks: list[str] = []

    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip():
            chunks.append(f"[paragraph {index + 1}] {paragraph.text.strip()}")

    for table_index, table in enumerate(document.tables):
        chunks.append(f"--- TABLE {table_index + 1} ---")
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            chunks.append(f"[table {table_index + 1} row {row_index + 1}] " + " | ".join(cells))

    return DocumentPayload(
        file=filename,
        text="\n".join(chunks),
        locator_hint="paragraph <n> / table <n> row <n>",
        reader="python-docx",
    )


# ---------------------------------------------------------------------------
# plain text
# ---------------------------------------------------------------------------

def read_text(handle, filename: str) -> DocumentPayload:
    raw = _decode(handle.read())
    numbered = "\n".join(f"[line {i + 1}] {line}" for i, line in enumerate(raw.splitlines()))
    return DocumentPayload(file=filename, text=numbered,
                           locator_hint="line <n>", reader="text")


# ---------------------------------------------------------------------------
# email
# ---------------------------------------------------------------------------

def read_eml(handle, filename: str, depth: int = 0) -> DocumentPayload:
    """Read a saved email properly, rather than as a wall of raw text.

    Treating an .eml as plain text mostly works and then fails exactly where it
    matters: a body sent as quoted-printable turns "₹42/kg" into "=E2=82=B942/kg",
    a base64 body reads as gibberish, and an attached quotation -- which is how
    most suppliers actually send a price list -- is invisible.

    So the message is parsed: headers first, because "From:" is often the only
    place a supplier's identity appears in a forwarded thread; then the plain
    body; then every attachment, opened with the same readers any other file
    would get. A PDF quotation attached to a covering email ends up read, with
    its own locators, in the same payload as the note that carried it.
    """
    import email
    from email import policy

    raw = handle.read()
    if isinstance(raw, str):
        raw = raw.encode("utf-8", errors="ignore")
    message = email.message_from_bytes(raw, policy=policy.default)

    # A forwarded thread carries the earlier message as an attachment, which is
    # read by this same function -- so a chain of forwards nests, and a file
    # built to nest into itself recurses until the interpreter runs out of
    # stack. Three levels covers "supplier replied to a forward of our request"
    # with room to spare; beyond that the content is summarised rather than
    # opened.
    if depth > MAX_EMAIL_DEPTH:
        return DocumentPayload(
            file=filename,
            text=_decode(raw)[:20000],
            locator_hint="line <n>",
            reader="email (nesting limit)",
            warnings=[f"'{filename}' is nested more than {MAX_EMAIL_DEPTH} "
                      f"messages deep; it was read as plain text rather than "
                      f"unpacked any further."],
        )

    chunks: list[str] = ["--- EMAIL HEADERS ---"]
    for header in ("From", "To", "Cc", "Date", "Sent", "Subject"):
        value = message.get(header)
        if value:
            chunks.append(f"[header] {header}: {value}")

    images: list[bytes] = []
    warnings: list[str] = []
    attachments: list[str] = []
    reader_names = ["email"]

    body = message.get_body(preferencelist=("plain", "html"))
    if body is not None:
        text = body.get_content()
        if body.get_content_subtype() == "html":
            text = re.sub(r"<br\s*/?>", "\n", text)
            text = re.sub(r"<[^>]+>", " ", text)
            warnings.append("The message had no plain-text part; the HTML body "
                            "was stripped to text.")
        chunks.append("--- EMAIL BODY ---")
        chunks.extend(f"[line {i + 1}] {line}"
                      for i, line in enumerate(text.splitlines()))

    for part in message.iter_attachments():
        name = part.get_filename() or "attachment"
        payload = part.get_payload(decode=True)
        if not payload:
            continue
        attachments.append(name)
        try:
            nested = load(io.BytesIO(payload), name, depth=depth + 1)
        except Exception as exc:
            warnings.append(f"Attachment '{name}' could not be read: {exc}")
            continue
        if nested.text:
            chunks.append(f"--- ATTACHMENT: {name} ---")
            chunks.append(nested.text)
        if nested.images:
            images.extend(nested.images)
        reader_names.append(nested.reader)
        warnings.extend(nested.warnings)

    if attachments:
        warnings.insert(0, "Attached to the email: " + ", ".join(attachments) + ".")

    return DocumentPayload(
        file=filename,
        text="\n".join(chunks),
        images=images,
        locator_hint="email line <n>, or 'attachment <name>' where one was read",
        reader=" + ".join(dict.fromkeys(reader_names)),
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# images
# ---------------------------------------------------------------------------

def read_image(handle, filename: str) -> DocumentPayload:
    blob = handle.read()
    ext = extension_of(filename)
    mime = ("image/png" if ext == "png" else
            "image/webp" if ext == "webp" else
            "image/jpeg" if ext in {"jpg", "jpeg"} else
            f"image/{ext}")
    warnings = ["Read by vision model directly from the image."]

    # A phone photo is often a .heic, and a scanner very often produces a
    # .tiff or a .bmp. Both were accepted at the upload box and then sent with
    # a jpeg content type, which the model rejects -- so the buyer got an API
    # error about MIME types for a file the application had told them it could
    # read. Convert instead, and only say we cannot when the converter is
    # genuinely missing.
    if mime not in VISION_MIMES:
        converted, note = _to_png(blob, ext)
        if converted is not None:
            blob, mime = converted, "image/png"
        warnings.append(note)

    return DocumentPayload(
        file=filename,
        images=[blob],
        image_mime=mime,
        locator_hint="photo",
        reader="vision",
        warnings=warnings,
    )


def _to_png(blob: bytes, ext: str) -> tuple[Optional[bytes], str]:
    """Re-encode an image the vision model will not take. (data, explanation)."""
    try:
        from PIL import Image
    except ImportError:
        return None, (f"A .{ext} image cannot be sent to the model directly and "
                      f"Pillow is not installed to convert it. Install Pillow, "
                      f"or re-save the photograph as a JPEG or PNG.")
    try:
        with Image.open(io.BytesIO(blob)) as image:
            if image.mode not in ("RGB", "L"):
                image = image.convert("RGB")
            out = io.BytesIO()
            image.save(out, format="PNG")
        return out.getvalue(), f"Converted from .{ext} to PNG so it could be read."
    except Exception as exc:
        return None, (f"A .{ext} image could not be converted for reading "
                      f"({type(exc).__name__}). Re-save it as a JPEG or PNG.")


# ---------------------------------------------------------------------------
# pdf, with a vision fallback
# ---------------------------------------------------------------------------

def _looks_corrupted(text: str) -> tuple[bool, str]:
    """Decide whether extracted PDF text can be trusted.

    Two failure modes matter: nothing came out at all (a scan), and glyphs came
    out as replacement boxes (a broken font map, which eats currency symbols).
    """
    if not text.strip():
        return True, "no extractable text layer (likely a scan)"

    marker_count = sum(text.count(marker) for marker in _CORRUPTION_MARKERS)
    if marker_count >= 3:
        return True, (f"{marker_count} unmapped glyphs in the text layer "
                      "(currency symbols are probably being lost)")

    # A price-bearing document with no currency symbol and no currency code at
    # all is suspicious enough to double-check with the vision model.
    has_digits = bool(re.search(r"\d", text))
    has_currency = bool(re.search(r"[₹$€£]|INR|USD|EUR|GBP", text, re.I))
    if has_digits and not has_currency and marker_count:
        return True, "prices present but no currency marker survived extraction"

    return False, ""


def _render_pdf_pages(data: bytes, max_pages: int = 12) -> list[bytes]:
    """Render pages to PNG so the vision model can read them."""
    try:
        import pymupdf  # PyMuPDF
    except ImportError:  # pragma: no cover
        try:
            import fitz as pymupdf  # older name
        except ImportError:
            return []

    images: list[bytes] = []
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        for page in doc[:max_pages]:
            pix = page.get_pixmap(dpi=170)
            images.append(pix.tobytes("png"))
    return images


def read_pdf(handle, filename: str) -> DocumentPayload:
    data = handle.read() if hasattr(handle, "read") else handle
    reader = PdfReader(io.BytesIO(data))

    chunks: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            chunks.append(f"--- PAGE {page_number} ---")
            chunks.append(page_text)

    text = "\n".join(chunks)
    corrupted, why = _looks_corrupted(text)

    payload = DocumentPayload(
        file=filename,
        text="" if corrupted else text,
        locator_hint="page <n>",
        reader="pypdf",
    )

    if corrupted:
        images = _render_pdf_pages(data)
        if images:
            payload.images = images
            payload.image_mime = "image/png"
            payload.reader = "pypdf -> vision fallback"
            payload.warnings.append(
                f"Text layer rejected ({why}); pages re-read as images."
            )
        else:
            # No renderer available: fall back to the flawed text but say so
            # loudly rather than silently reporting corrupted prices.
            payload.text = text
            payload.reader = "pypdf (degraded)"
            payload.warnings.append(
                f"Text layer is unreliable ({why}) and PyMuPDF is not installed, "
                "so the page could not be re-read as an image. Prices from this "
                "document should be treated as unverified."
            )

    return payload


# ---------------------------------------------------------------------------
# dispatcher
# ---------------------------------------------------------------------------

def load(handle, filename: str, depth: int = 0) -> DocumentPayload:
    """Read any supported vendor file into a payload."""
    ext = extension_of(filename)

    if ext == "pdf":
        return read_pdf(handle, filename)
    if ext in {"xlsx", "xlsm", "xls"}:
        return read_spreadsheet(handle, filename)
    if ext == "docx":
        return read_docx(handle, filename)
    if ext == "csv":
        return read_csv(handle, filename)
    if ext == "eml":
        return read_eml(handle, filename, depth=depth)
    if ext == "msg":
        # An Outlook .msg is a compound OLE file, not an RFC-822 message. It
        # was being handed to the email parser, which found no headers and no
        # body and returned a payload containing the words "EMAIL HEADERS" and
        # nothing else -- so the supplier appeared to have sent an empty
        # quotation. Say what happened instead.
        return DocumentPayload(
            file=filename,
            reader="unsupported",
            warnings=["Outlook .msg is a different format from .eml and cannot "
                      "be read here. Drag the message out of Outlook onto your "
                      "desktop to get an .eml, or upload the supplier's "
                      "attachment on its own."],
        )
    if ext in TEXT_EXTENSIONS:
        return read_text(handle, filename)
    if ext in IMAGE_EXTENSIONS:
        return read_image(handle, filename)

    return DocumentPayload(
        file=filename,
        reader="unsupported",
        warnings=[f"Unsupported file type '.{ext}'."],
    )


def load_path(path: str) -> DocumentPayload:
    with open(path, "rb") as handle:
        return load(handle, os.path.basename(path))


SUPPORTED_UPLOAD_TYPES = sorted(
    {"pdf", "xlsx", "xls", "xlsm", "docx", "txt", "csv", "md", "eml"} | IMAGE_EXTENSIONS
)
